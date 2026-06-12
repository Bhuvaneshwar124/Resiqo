import re
import io
import pymupdf
from docx import Document

class DocumentParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        text = []
        try:
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text.append(page.get_text())
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
        
        return "\n".join(text)

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        text = []
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract text from tables as well
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        text.append(" | ".join(row_data))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
            
        return "\n".join(text)

    @staticmethod
    def parse_document(file_bytes: bytes, filename: str) -> str:
        """Extracts raw text from a PDF or DOCX file."""
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            return DocumentParser.extract_text_from_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            return DocumentParser.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF and DOCX are supported.")

document_parser = DocumentParser()
